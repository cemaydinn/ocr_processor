import os
import logging
from typing import List, Union
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
import multiprocessing as mp
from concurrent.futures import ThreadPoolExecutor


class OCRProcessor:
    def __init__(self):
        self.supported_languages = ['eng', 'tur']  # dil desteği
        self.logger = self._setup_logger()

    def _setup_logger(self):
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        return logging.getLogger(__name__)

    def process_document(self, file_path: str) -> dict:
        """
        farklı uzantıları işleme ve metin çıkarma
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()

            if file_extension in ['.pdf']:
                return self._process_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._process_word(file_path)
            elif file_extension in ['.png', '.jpg', '.jpeg']:
                return self._process_image(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")

        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            raise

    def _process_pdf(self, file_path: str) -> dict:
        """PDF işleme"""
        images = convert_from_path(file_path)
        with ThreadPoolExecutor() as executor:
            results = list(executor.map(self._extract_text_from_image, images))

        return {
            'text': '\n'.join(results),
            'pages': len(images),
            'format': 'pdf'
        }

    def _process_word(self, file_path: str) -> dict:
        """docx dosyasını işleme"""
        doc = Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        return {
            'text': text,
            'pages': len(doc.paragraphs),
            'format': 'docx'
        }

    def _process_image(self, file_path: str) -> dict:
        """image dosya işleme"""
        image = Image.open(file_path)
        text = self._extract_text_from_image(image)

        return {
            'text': text,
            'pages': 1,
            'format': 'image'
        }

    def _extract_text_from_image(self, image: Image) -> str:
        """OCR ile görüntüden metin çıkarma"""
        try:
            text = pytesseract.image_to_string(
                image,
                lang='+'.join(self.supported_languages)
            )
            return text.strip()
        except Exception as e:
            self.logger.error(f"OCR extraction error: {str(e)}")
            raise



from .utils import (
            setup_logging,
            validate_file,
            clean_text,
            get_file_info,
            create_output_filename
        )